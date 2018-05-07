from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor



condons2AminoAcid = {
    'GCT':'A', 'GCC':'A', 'GCA':'A', 'GCG':'A',
    'CGT':'R', 'CGC':'R', 'CGA':'R', 'CGG':'R', 'AGA':'R', 'AGG':'R',
    'AAT':'N', 'AAC':'N',
    'GAT':'D', 'GAC':'D',
    'TGT':'C', 'TGC':'C',
    'CAA':'Q', 'CAG':'Q',
    'GAA':'E', 'GAG':'E',
    'GGT':'G', 'GGC':'G', 'GGA':'G', 'GGG':'G',
    'CAT':'H', 'CAC':'H',
    'ATT':'I', 'ATC':'I', 'ATA':'I',
    'TTA':'L', 'TTG':'L', 'CTT':'L', 'CTC':'L', 'CTA':'L', 'CTG':'L',
    'AAA':'K', 'AAG':'K',
    'ATG':'M',
    'TTT':'F', 'TTC':'F',
    'CCT':'P', 'CCC':'P', 'CCA':'P', 'CCG':'P',
    'TCT':'S', 'TCC':'S', 'TCA':'S', 'TCG':'S', 'AGT':'S', 'AGC':'S',
    'ACT':'T', 'ACC':'T', 'ACA':'T', 'ACG':'T',
    'TGG':'W',
    'TAT':'Y', 'TAC':'Y',
    'GTT':'V', 'GTC':'V', 'GTA':'V', 'GTG':'V'
}

aminoAcid2Color = {
    'A':(0x01,0xCD,0xB4),
    'R':(0x00,0x00,0xFF),
    'N':(0x80,0x80,0xC0),
    'D':(0xFF,0x00,0x00),
    'C':(0xA2,0x00,0x00),
    'Q':(0x71,0x71,0xB9),
    'E':(0xFF,0x00,0x00),
    'G':(0xE7,0x96,0x01),
    'H':(0xEE,0x2D,0xEE),
    'I':(0x00,0x80,0x40),
    'L':(0x00,0x80,0x40),
    'K':(0x00,0x00,0xD9),
    'M':(0x00,0x80,0x40),
    'F':(0x15,0xBF,0xAE),
    'P':(0x7E,0x73,0x63),
    'S':(0x53,0x7E,0xAA),
    'T':(0x53,0x7E,0xAA),
    'W':(0x27,0xAB,0x60),
    'Y':(0x15,0xBD,0xAC),
    'V':(0x00,0x80,0x40)
}


base2Color = {
    'T':(0xFF,0x00,0x00),
    'C':(0x00,0x00,0xFF),
    'A':(0x00,0x80,0x00),
    'G':(0x00,0x00,0x00),
}


testinstanfce = 'GAGGTGCAGCTGGTCGAGAGCGGCGGAGGGCTGGTGAAGGCCGGAGGAAGCCTGATCCTGAGCTGCGGCGTGAGCAACTTCAGGATCAGCGCCCACACCATGAACTGGGTGCGGAGGGTGCCAGGCGGCGGACTGGAGTGGGTGGCCAGCATCAGCACCAGCAGCACCTACAGGGACTAC'
expectedResult =  'EVQLVESGGGLVKAGGSLIL'

testName1 = '2G12_vH'
testName2 = '353/11_vH'

def getNumOfSpaces(number):
    if len(number) is 2:
        return 8
    elif len(number) is 3:
        return 7
    else:
        return None

def convertBasesSequence2AminoAcid(basesSequence):
    stepSize = 3
    condons = [basesSequence[i:i + stepSize] for i in range(0, len(basesSequence), stepSize)]
    aminoAcids = [condons2AminoAcid[condo] for condo in condons]

    return aminoAcids

def printBases(paragraph, basesSequence):
    for base in basesSequence:
        run = paragraph.add_run(base)
        run.bold = True
        font = run.font
        baseColor = base2Color[base]
        font.color.rgb = RGBColor(baseColor[0], baseColor[1], baseColor[2])
    paragraph.add_run().add_break()

def printNumbering(p, start, end):
    numbering = ["{:d}".format(x) for x in range(start, end+1, 10)]
    numberingSting = [(' ' * getNumOfSpaces(number)) + number for number in numbering]
    p.add_run(' ' + ''.join(numberingSting))

def printIndicator(paragraph):
    indicator = '....|....|....|....|....|....|....|....|....|....|....|....|'
    paragraph.add_run(indicator).add_break()


def printAminoAcids(paragraph, baseSequence):
    aminoAcids = convertBasesSequence2AminoAcid(baseSequence)

    print(''.join(aminoAcids) == expectedResult)

    numberOfSpaces = 1
    for aminoAcid in aminoAcids:
        run = paragraph.add_run((' '*numberOfSpaces) + aminoAcid)
        run.bold = True
        font = run.font
        aminoAcidColor = aminoAcid2Color[aminoAcid]
        font.color.rgb = RGBColor(aminoAcidColor[0], aminoAcidColor[1], aminoAcidColor[2])

        numberOfSpaces = 2
    paragraph.add_run().add_break()

def printUnknown(paragraph,baseSequence):
    paragraph.add_run().add_break()

def printStars(paragraph,baseSequence):
    aminoAcids = convertBasesSequence2AminoAcid(baseSequence)

    print(''.join(aminoAcids) == expectedResult)

    numberOfSpaces = 1
    for aminoAcid in aminoAcids:
        run = paragraph.add_run((' '*numberOfSpaces) + '*')
        run.bold = True
        font = run.font
        aminoAcidColor = aminoAcid2Color[aminoAcid]
        font.color.rgb = RGBColor(aminoAcidColor[0], aminoAcidColor[1], aminoAcidColor[2])

        numberOfSpaces = 2

def generateAndPersistDocument(filename, names, sequencce):
    name1 = names[0]
    name2 = names[1]

    stepSize = 60
    sequencesPerLines = [sequencce[i:i + stepSize] for i in range(0, len(sequencce), stepSize)]

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(10)

    for index,sequencesPerLine in enumerate(sequencesPerLines):
        startPosition = 60 * index + 10
        endPosition = startPosition + 50
        p = document.add_paragraph()

        p.add_run(' '*11)

        printNumbering(p,startPosition, endPosition)
        p = document.add_paragraph()

        p.add_run(' ' * 11)
        printIndicator(p)

        p.add_run(' ' * 11)
        printAminoAcids(p,sequencesPerLine)

        p.add_run(name1 + ' ' * (11 - len(name1))).bold = True
        printBases(p, sequencesPerLine)

        p.add_run(name2 + ' ' * (11 - len(name2))).bold = True
        printUnknown(p,sequencesPerLine)

        p.add_run(' ' * 11)
        printStars(p,sequencesPerLine)

    document.save(filename)


'''
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

run = document.add_paragraph('karlsplatz ').add_run('is supa')
font = run.font
font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
p=document.add_paragraph('aaa')
'''